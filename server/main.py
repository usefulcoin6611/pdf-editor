#!/usr/bin/env python3
"""
PDF Editor Backend Server (V4 - Engine Logging & Diagnostics)
"""

import os
import uuid
import shutil
import subprocess
import tempfile
import logging
from pathlib import Path
from typing import Dict, List, Optional, Any

import uvicorn
from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel

# ─── Logging ──────────────────────────────────────────────────────────────────
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger("pdf-backend")

# ─── App ──────────────────────────────────────────────────────────────────────

app = FastAPI(title="PDF Editor Backend", version="4.0.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["X-Applied-Edits", "X-Engine"]
)

TEMP_DIR = Path(tempfile.gettempdir()) / "pdf_editor_sessions"
TEMP_DIR.mkdir(exist_ok=True)

SESSIONS: Dict[str, Dict[str, Path]] = {}

# ─── Models ───────────────────────────────────────────────────────────────────

class RunInfo(BaseModel):
    id: str
    text: str
    bold: bool
    italic: bool
    underline: bool
    fontSize: Optional[float]
    fontName: Optional[str]
    color: Optional[str]

class ParagraphInfo(BaseModel):
    id: str
    text: str
    alignment: Optional[str]
    lineSpacing: Optional[float]
    runs: List[RunInfo]

class LoadResponse(BaseModel):
    sessionId: str
    paragraphs: List[ParagraphInfo]
    pageCount: int

class EditRun(BaseModel):
    runId: str
    newText: str
    fontSize: Optional[float] = None

class ExportRequest(BaseModel):
    sessionId: str
    edits: List[EditRun]

class SplitRange(BaseModel):
    label: str
    pages: str # e.g. "1-3", "5", "7-10"

# ─── Core Logic ───────────────────────────────────────────────────────────────

def get_color_hex(run) -> Optional[str]:
    try:
        color = run.font.color
        if color and color.rgb:
            return str(color.rgb)
        return None
    except Exception:
        return None

def parse_docx(docx_path: Path) -> List[Dict[str, Any]]:
    from docx import Document
    doc = Document(str(docx_path))
    paragraphs = []

    for p_idx, para in enumerate(doc.paragraphs):
        if not para.text.strip(): continue

        fmt = para.paragraph_format
        align_map = {0: "left", 1: "center", 2: "right", 3: "justify"}
        align_str = align_map.get(para.alignment, "left") if para.alignment is not None else "left"
        line_spacing = fmt.line_spacing if fmt.line_spacing is not None else 1.0

        runs = []
        for r_idx, run in enumerate(para.runs):
            if not run.text: continue

            font_size = run.font.size.pt if run.font.size else (para.style.font.size.pt if para.style and para.style.font and para.style.font.size else None)
            font_name = run.font.name or (para.style.font.name if para.style and para.style.font else None)

            runs.append({
                "id": f"{p_idx}-{r_idx}",
                "text": run.text,
                "bold": bool(run.bold),
                "italic": bool(run.italic),
                "underline": bool(run.underline),
                "fontSize": font_size,
                "fontName": font_name,
                "color": get_color_hex(run),
            })

        if runs:
            paragraphs.append({
                "id": str(p_idx),
                "text": para.text,
                "alignment": align_str,
                "lineSpacing": line_spacing,
                "runs": runs,
            })
    return paragraphs


def apply_edits(src_docx: Path, edits: List[EditRun], dst_docx: Path):
    from docx import Document
    from docx.shared import Pt
    doc = Document(str(src_docx))
    edit_map = {e.runId: e for e in edits}
    applied_count = 0
    requested_ids = set(edit_map.keys())
    found_ids = set()

    for p_idx, para in enumerate(doc.paragraphs):
        # Paragraph-level fallback: if a user sends a p_idx directly
        if str(p_idx) in edit_map:
            # Re-build the paragraph with new text
            edit_obj = edit_map[str(p_idx)]
            new_text = edit_obj.newText
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = new_text
                if edit_obj.fontSize is not None:
                    para.runs[0].font.size = Pt(edit_obj.fontSize * 0.75)
            else:
                para.add_run(new_text)
            
            found_ids.add(str(p_idx))
            applied_count += 1
            logger.info(f"Applied paragraph-level edit to para {p_idx}")
            continue

        fmt = para.paragraph_format
        s_line = fmt.line_spacing
        s_rule = fmt.line_spacing_rule
        s_before = fmt.space_before
        s_after = fmt.space_after

        modified = False
        for r_idx, run in enumerate(para.runs):
            rid = f"{p_idx}-{r_idx}"
            if rid in edit_map:
                edit_obj = edit_map[rid]
                run.text = edit_obj.newText
                if edit_obj.fontSize is not None:
                    run.font.size = Pt(edit_obj.fontSize * 0.75)
                modified = True
                found_ids.add(rid)
                applied_count += 1
        
        if modified:
            fmt.line_spacing = s_line
            fmt.line_spacing_rule = s_rule
            fmt.space_before = s_before
            fmt.space_after = s_after
            logger.info(f"Applied run-level edit to para {p_idx}, locked spacing")

    missing = requested_ids - found_ids
    if missing:
        logger.warning(f"Some edits were NOT applied (IDs not found): {missing}")
    
    logger.info(f"Total edits applied: {applied_count}/{len(edits)}")
    doc.save(str(dst_docx))
    return applied_count


def docx_to_pdf(docx_path: Path, pdf_path: Path) -> str:
    """Returns the name of the engine used, or 'failed'."""
    # Priority 1: MS Word via docx2pdf
    if os.name == 'nt':
        try:
            from docx2pdf import convert
            convert(str(docx_path), str(pdf_path))
            if pdf_path.exists(): 
                logger.info("Export success via MS Word (docx2pdf)")
                return "ms-word"
        except Exception as e:
            logger.warning(f"MS Word export failed: {e}")

    # Priority 2: LibreOffice
    lo_candidates = ["soffice", "libreoffice", r"C:\Program Files\LibreOffice\program\soffice.exe"]
    for cmd in lo_candidates:
        try:
            subprocess.run([cmd, "--headless", "--convert-to", "pdf", "--outdir", str(pdf_path.parent), str(docx_path)], timeout=60, capture_output=True)
            expected = pdf_path.parent / (docx_path.stem + ".pdf")
            if expected.exists():
                if expected != pdf_path: shutil.move(str(expected), str(pdf_path))
                logger.info(f"Export success via LibreOffice ({cmd})")
                return "libreoffice"
        except Exception: continue

    return "failed"

# ─── Endpoints ────────────────────────────────────────────────────────────────

@app.get("/api/health")
def health():
    return {
        "status": "ok",
        "has_docx2pdf": True,
        "os": os.name
    }

@app.post("/api/load-pdf", response_model=LoadResponse)
async def load_pdf(file: UploadFile = File(...)):
    session_id = str(uuid.uuid4())
    session_dir = TEMP_DIR / session_id
    session_dir.mkdir(parents=True, exist_ok=True)

    pdf_path = session_dir / "original.pdf"
    pdf_path.write_bytes(await file.read())

    docx_path = session_dir / "document.docx"
    try:
        from pdf2docx import Converter
        cv = Converter(str(pdf_path))
        cv.convert(str(docx_path), start=0, end=None)
        cv.close()
    except Exception as e:
        logger.error(f"Conversion error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

    SESSIONS[session_id] = {"pdf": pdf_path, "docx": docx_path, "dir": session_dir}
    return LoadResponse(sessionId=session_id, paragraphs=parse_docx(docx_path), pageCount=0)

@app.post("/api/export")
async def export_pdf(request: ExportRequest):
    session = SESSIONS.get(request.sessionId)
    if not session: raise HTTPException(status_code=404, detail="Session expired")

    dst_docx = session["dir"] / "edited.docx"
    out_pdf = session["dir"] / "final.pdf"
    
    # Ensure no stale output exists
    if out_pdf.exists(): out_pdf.unlink()

    logger.info(f"Received export request for session {request.sessionId} with {len(request.edits)} edits")
    for e in request.edits:
        logger.info(f"  - Edit: ID={e.runId}, TextLen={len(e.newText)}")

    applied_count = apply_edits(session["docx"], request.edits, dst_docx)
    
    # Store for debug access
    session["edited_docx"] = dst_docx

    engine = docx_to_pdf(dst_docx, out_pdf)
    if engine != "failed":
        return FileResponse(
            path=out_pdf, 
            filename="edited.pdf", 
            headers={
                "X-Engine": engine,
                "X-Applied-Edits": str(applied_count)
            }
        )
    raise HTTPException(status_code=500, detail="PDF conversion failed on server")

@app.post("/api/split-pdf")
async def split_pdf(file: UploadFile = File(...), ranges: str = Form("all")):
    """
    Splits PDF into multiple parts based on ranges.
    ranges is a JSON string of List[Dict[label: str, pages: str]]
    """
    import json
    import fitz
    import zipfile
    import io

    try:
        range_data = json.loads(ranges)
    except Exception:
        range_data = [{"label": "part1", "pages": "1-end"}]

    pdf_bytes = await file.read()
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    total_pages = doc.page_count

    from fastapi import Response
    # If only one range and it's "all", just return the file
    if len(range_data) == 1 and (range_data[0]["pages"] == "all" or range_data[0]["pages"] == f"1-{total_pages}"):
        return Response(content=pdf_bytes, media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=split.pdf"})

    zip_buffer = io.BytesIO()
    try:
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            for item in range_data:
                label = item.get("label", "part")
                pages_str = item.get("pages", "")
                
                output_doc = fitz.open()
                try:
                    page_indices = []
                    for part in pages_str.split(","):
                        part = part.strip()
                        if not part: continue
                        if "-" in part:
                            try:
                                start_str, end_str = part.split("-")
                                start_idx = int(start_str) - 1
                                if end_str.lower() == "end":
                                    end_idx = total_pages - 1
                                else:
                                    end_idx = int(end_str) - 1
                                page_indices.extend(range(start_idx, end_idx + 1))
                            except ValueError: continue
                        else:
                            try:
                                page_indices.append(int(part) - 1)
                            except ValueError: continue
                    
                    page_indices = [i for i in page_indices if 0 <= i < total_pages]
                    if not page_indices:
                        continue

                    output_doc.insert_pdf(doc, from_page=0, to_page=total_pages-1, select=page_indices)
                    
                    # Use tobytes() for reliable in-memory conversion
                    pdf_data = output_doc.tobytes()
                    zip_file.writestr(f"{label}.pdf", pdf_data)
                    output_doc.close()
                except Exception as e:
                    logger.error(f"Error processing range {pages_str}: {e}")
                    if not output_doc.is_closed: output_doc.close()
                    continue
    finally:
        doc.close()

    from fastapi import Response
    return Response(
        content=zip_buffer.getvalue(), 
        media_type="application/zip",
        headers={"Content-Disposition": "attachment; filename=split_pdfs.zip"}
    )

@app.post("/api/export-docx")
async def export_docx(request: ExportRequest):
    session = SESSIONS.get(request.sessionId)
    if not session: raise HTTPException(status_code=404, detail="Session expired")

    dst_docx = session["dir"] / "edited.docx"
    
    logger.info(f"Received DOCX export request for session {request.sessionId} with {len(request.edits)} edits")
    applied_count = apply_edits(session["docx"], request.edits, dst_docx)
    
    return FileResponse(
        path=dst_docx, 
        filename="edited.docx", 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "X-Applied-Edits": str(applied_count)
        }
    )

@app.get("/api/debug-docx/{session_id}")
async def debug_docx(session_id: str):
    session = SESSIONS.get(session_id)
    if not session or "edited_docx" not in session:
        raise HTTPException(status_code=404, detail="Debug file not found")
    return FileResponse(path=session["edited_docx"], filename="debug_edited.docx")

@app.post("/api/merge-pdf")
async def merge_pdf(files: List[UploadFile] = File(...)):
    """
    Merges multiple PDF files into one.
    """
    import fitz
    import io
    from fastapi import Response

    if not files:
        raise HTTPException(status_code=400, detail="No files provided")

    merged_doc = fitz.open()
    try:
        for file in files:
            pdf_bytes = await file.read()
            src_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            merged_doc.insert_pdf(src_doc)
            src_doc.close()
        
        pdf_data = merged_doc.tobytes()
        return Response(
            content=pdf_data,
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=merged.pdf"}
        )
    except Exception as e:
        logger.error(f"Merge error: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        merged_doc.close()

if __name__ == "__main__":
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
